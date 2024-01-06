import pandas as pd
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx2pdf import convert
import os
import datetime
import tkinter as tk
from tkinter import messagebox
import webbrowser
from tkcalendar import DateEntry
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import gspread
from oauth2client.service_account import ServiceAccountCredentials

def fetch_reference_number_from_google_sheets(name):    
    #spreadsheet_id = 'https://docs.google.com/spreadsheets/d/1oxKGdu0fkyCZuBW9oUAuk_IT41lMyNM_HJSuzKPJFS0/edit#gid=0'
    spreadsheet_id = '1oxKGdu0fkyCZuBW9oUAuk_IT41lMyNM_HJSuzKPJFS0'
    credentials_file = 'internal-project-395406-815b0a5abb4b.json'
    try:
        scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
        credentials = ServiceAccountCredentials.from_json_keyfile_name(credentials_file, scope)
        client = gspread.authorize(credentials)

        spreadsheet = client.open_by_key(spreadsheet_id)
        worksheet = spreadsheet.get_worksheet(0)  
        records = worksheet.get_all_records()
        df = pd.DataFrame.from_records(records)
        reference_number = df.loc[df['Name'] == name, 'Reference Number'].values[0]
        return reference_number if not pd.isnull(reference_number) else None
    except Exception as e:
        print("Error occurred while fetching reference number:", e)
        return None

def replace_placeholder_with_formatting(run, placeholder, value, is_bold=False):
    parts = run.text.split(placeholder)
    run.clear()
    for part in parts:
        run.add_text(part)
        if part == "":
            new_run = run.add_run(value)
            font = new_run.font
            font.bold = is_bold
            
def create_pdf_from_template(name, usn, guide, start_date, college_name, domain, gender, tdate, output_format):
    if gender == "Male":
        template_path = "Internship Persuing Template_m.docx"
    elif gender == "Female":
        template_path = "Internship Persuing Template_fe.docx"
    else:
        messagebox.showerror("Error", "Invalid gender selected.")
        return

    reference_number = fetch_reference_number_from_google_sheets(name)
    if reference_number is None:
        messagebox.showerror("Error", f"No reference number found for {name}.")
        return

    doc = Document(template_path)
    # Replace placeholders in the template
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            font = run.font
            font.size = Pt(12)
            
            if "{{name}}" in run.text:
                replace_placeholder_with_formatting(run, "{{name}}", name, True)
            elif "{{usn}}" in run.text:
                replace_placeholder_with_formatting(run, "{{usn}}", usn, True)
            elif "{{college_name}}" in run.text:
                run.text = run.text.replace("{{college_name}}", college_name)
            elif "{{start_date}}" in run.text:
                run.text = run.text.replace("{{start_date}}", start_date)
            elif "{{domain}}" in run.text:
                run.text = run.text.replace("{{domain}}", domain)
            elif "{{guide}}" in run.text:
                run.text = run.text.replace("{{guide}}", guide)
            elif "{{tdate}}" in run.text:
                run.text = run.text.replace("{{tdate}}", tdate)
            elif "{{reference_number}}" in run.text:
                run.text = run.text.replace("{{reference_number}}", str(reference_number))

    current_date = datetime.datetime.now().strftime("%Y%m%d")
    output_filename = f"{name}_{current_date}.{output_format}"
    doc.save('temp.docx')

    if output_format == "pdf":
        convert("temp.docx", output_filename)
    else:
        os.rename('temp.docx', output_filename)

    messagebox.showinfo("File Created", f"File created successfully: {output_filename}")
    show_link(output_filename)

def show_link(filepath):
    link_label.config(text="File created: " + filepath)
    open_button.config(state=tk.NORMAL, command=lambda: open_file(filepath))

def open_file(filepath):
    webbrowser.open(filepath)

def reset_inputs():
    name_entry.delete(0, tk.END)
    usn_entry.delete(0, tk.END)
    guide_entry.delete(0, tk.END)
    start_date_entry.delete(0, tk.END)
    college_name_entry.delete(0, tk.END)
    domain_entry.delete(0, tk.END)

def create_pdf():
    global tdate
    name = name_entry.get()
    usn = usn_entry.get()
    guide = guide_entry.get()
    start_date = start_date_entry.get()
    college_name = college_name_entry.get()
    domain = domain_entry.get()
    gender = gender_var.get()
    tdate = datetime.datetime.now().strftime("%B %d, %Y")

    if len(name) == 0:
        messagebox.showerror("Error", "Please enter a name.")
        return
    if len(usn) == 0:
        messagebox.showerror("Error", "Please enter a USN.")
        return
    if len(guide) == 0:
        messagebox.showerror("Error", "Please enter a guide name.")
        return
    if len(start_date) == 0:
        messagebox.showerror("Error", "Please enter a start date.")
        return
    if len(college_name) == 0:
        messagebox.showerror("Error", "Please enter a college name.")
        return
    if len(domain) == 0:
        messagebox.showerror("Error", "Please enter a domain name.")
        return

    create_pdf_from_template(name, usn, guide, start_date, college_name, domain, gender, tdate, "pdf")

# Rest of the code remains unchanged
def save_as_pdf():
    global tdate
    name = name_entry.get()
    usn = usn_entry.get()
    guide = guide_entry.get()
    start_date = start_date_entry.get()
    college_name = college_name_entry.get()
    domain = domain_entry.get()
    gender = gender_var.get()
    tdate = datetime.datetime.now().strftime("%B %d, %Y")

    if len(name) == 0:
        messagebox.showerror("Error", "Please enter a name.")
        return
    if len(usn) == 0:
        messagebox.showerror("Error", "Please enter a USN.")
        return
    if len(guide) == 0:
        messagebox.showerror("Error", "Please enter a guide name.")
        return
    if len(start_date) == 0:
        messagebox.showerror("Error", "Please enter a start date.")
        return
    if len(college_name) == 0:
        messagebox.showerror("Error", "Please enter a college name.")
        return
    if len(domain) == 0:
        messagebox.showerror("Error", "Please enter a domain name.")
        return

    create_pdf_from_template(name, usn, guide, start_date, college_name, domain, gender, tdate, "pdf")

def save_as_docx():
    global tdate
    name = name_entry.get()
    usn = usn_entry.get()
    guide = guide_entry.get()
    start_date = start_date_entry.get()
    college_name = college_name_entry.get()
    domain = domain_entry.get()
    gender = gender_var.get()
    tdate = datetime.datetime.now().strftime("%B %d, %Y")

    if len(name) == 0:
        messagebox.showerror("Error", "Please enter a name.")
        return
    if len(usn) == 0:
        messagebox.showerror("Error", "Please enter a USN.")
        return
    if len(guide) == 0:
        messagebox.showerror("Error", "Please enter a guide name.")
        return
    if len(start_date) == 0:
        messagebox.showerror("Error", "Please enter a start date.")
        return
    if len(college_name) == 0:
        messagebox.showerror("Error", "Please enter a college name.")
        return
    if len(domain) == 0:
        messagebox.showerror("Error", "Please enter a domain name.")
        return

    create_pdf_from_template(name, usn, guide, start_date, college_name, domain, gender, tdate, "docx")

def send_email():
    email = email_entry.get()
    if not email:
        messagebox.showerror("Error", "Please enter an email address.")
        return

    generated_file = link_label.cget("text")[14:]
    if not os.path.exists(generated_file):
        messagebox.showerror("Error", "The generated file does not exist.")
        return

    # Email configurations
    sender_email = "priyashingri107@gmail.com" 
    sender_password = "fzmfwnhjhjtuaouf"
    receiver_email = email

    msg = MIMEMultipart()
    msg["From"] = sender_email
    msg["To"] = receiver_email
    msg["Subject"] = "Internship Certificate from HanSoftUmans"

    with open(generated_file, "rb") as attachment:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(attachment.read())
        encoders.encode_base64(part)
        part.add_header(
            "Content-Disposition",
            f"attachment; filename= {generated_file}",
        )
        msg.attach(part)

    try:
        server = smtplib.SMTP("smtp.gmail.com", 587)
        server.starttls()
        server.login(sender_email, sender_password)
        server.sendmail(sender_email, receiver_email, msg.as_string())
        server.quit()
        messagebox.showinfo("Email Sent", "Email sent successfully.")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to send email: {e}")

# main window
window = tk.Tk()
window.title("HanSoftUmans")
window.geometry("700x800") 
window.resizable(True, True)  

# center alignment
screen_width = window.winfo_screenwidth()
screen_height = window.winfo_screenheight()
window_width = 400
window_height = 500
x = (screen_width - window_width) // 2
y = (screen_height - window_height) // 2
window.geometry(f"{window_width}x{window_height}+{x}+{y}")  # Center the window on the screen

# headlinename_entry
headline_label = tk.Label(window, text="HanSoftUmans", font=("Arial", 26, "bold"),fg="blue")
headline_label.grid(row=0, column=0, columnspan=2, pady=10)
subheadline_label = tk.Label(window, text="Internship Certificate", font=("Arial", 16, "bold"),fg="black")
subheadline_label.grid(row=1, column=0, columnspan=2, pady=10)

# labels and entry fields
name_label = tk.Label(window, text="Full Name:")
name_label.grid(row=2, column=0, padx=10, pady=5, sticky=tk.E)
name_entry = tk.Entry(window)
name_entry.grid(row=2, column=1, padx=10, pady=5, sticky=tk.W)
name_entry.focus()

usn_label = tk.Label(window, text="USN:")
usn_label.grid(row=3, column=0, padx=10, pady=5, sticky=tk.E)
usn_entry = tk.Entry(window)
usn_entry.grid(row=3, column=1, padx=10, pady=5, sticky=tk.W)

college_name_label = tk.Label(window, text="College Name:")
college_name_label.grid(row=4, column=0, padx=10, pady=5, sticky=tk.E)
college_name_entry = tk.Entry(window)
college_name_entry.grid(row=4, column=1, padx=10, pady=5, sticky=tk.W)

start_date_label = tk.Label(window, text="Start Date:")
start_date_label.grid(row=5, column=0, padx=10, pady=5, sticky=tk.E)
start_date_entry = DateEntry(window, width=12, background='black', foreground='white', date_pattern='dd/mm/yyyy')
start_date_entry.grid(row=5, column=1, padx=10, pady=5, sticky=tk.W)

guide_label = tk.Label(window, text="Guide:")
guide_label.grid(row=6, column=0, padx=10, pady=5, sticky=tk.E)
guide_entry = tk.Entry(window)
guide_entry.grid(row=6, column=1, padx=10, pady=5, sticky=tk.W)


domain_label = tk.Label(window, text="domain:")
domain_label.grid(row=7, column=0, padx=10, pady=5, sticky=tk.E)
domain_entry = tk.Entry(window)
domain_entry.grid(row=7, column=1, padx=10, pady=5, sticky=tk.W)

# variable
gender_var = tk.StringVar()
gender_var.set("None") 

# radio button
gender_label = tk.Label(window, text="Gender:")
gender_label.grid(row=8, column=0, padx=10, pady=5, sticky=tk.E)
male_radio = tk.Radiobutton(window, text="Male", variable=gender_var, value="Male")
male_radio.grid(row=8, column=1, padx=10, pady=5, sticky=tk.W)
female_radio = tk.Radiobutton(window, text="Female", variable=gender_var, value="Female")
female_radio.grid(row=9, column=1, padx=10, pady=5, sticky=tk.W)

email_label = tk.Label(window, text="Email:")
email_label.grid(row=12, column=0, padx=10, pady=5, sticky=tk.E)
email_entry = tk.Entry(window)
email_entry.grid(row=12, column=1, padx=10, pady=5, sticky=tk.W)

save_pdf_button = tk.Button(window, text="Save as PDF", command=save_as_pdf)
save_pdf_button.grid(row=10, column=0, padx=10, pady=10)

save_docx_button = tk.Button(window, text="Save as DOCX", command=save_as_docx)
save_docx_button.grid(row=10, column=1, padx=10, pady=10)

# file name
link_label = tk.Label(window, text="")
link_label.grid(row=10, column=0, columnspan=2, padx=10, pady=5)

open_button = tk.Button(window, text="Open File", state=tk.DISABLED)
open_button.grid(row=11, column=0, columnspan=2, padx=10, pady=5)
# send email button
send_email_button = tk.Button(window, text="Send Email", command=send_email)
send_email_button.grid(row=13, column=0, columnspan=2, padx=10, pady=10)
window.grid_rowconfigure(10, weight=1)  
window.grid_columnconfigure(0, weight=1)  
window.grid_columnconfigure(1, weight=1)  
window.mainloop()