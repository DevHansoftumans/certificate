from docx import Document
from docx.shared import Pt
from docx2pdf import convert
import os
import datetime
import tkinter as tk
from tkinter import messagebox
import webbrowser
from tkcalendar import DateEntry
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders

def create_pdf_from_template(name, usn, guide, start_date, college_name, domain, gender, tdate, output_format):
    if gender == "Male":
        template_path = "Internship Persuing Template_m.docx"
    elif gender == "Female":
        template_path = "Internship Persuing Template_fe.docx"
    else:
        messagebox.showerror("Error", "Invalid gender selected.")
        return

    doc = Document(template_path)

    # Replace placeholders in the template
    for paragraph in doc.paragraphs:
        if "{{name}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{name}}", name)
        if "{{usn}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{usn}}", usn)
        if "{{college_name}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{college_name}}", college_name)
        if "{{start_date}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{start_date}}", start_date)
        if "{{domain}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{domain}}", domain)
        if "{{guide}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{guide}}", guide)
        if "{{tdate}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{tdate}}", tdate)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            font = run.font
            font.size = Pt(12)

    current_date = datetime.datetime.now().strftime("%Y%m%d")
    output_filename = f"{name}_{current_date}.{output_format}"
    doc.save('temp.docx')

    if output_format == "pdf":
        convert("temp.docx", output_filename)
    else:
        os.rename('temp.docx', output_filename)

    messagebox.showinfo("File Created", f"File created successfully: {output_filename}")
    show_link(output_filename)

def show_link(filepath):
    link_label.config(text="File created: " + filepath)
    open_button.config(state=tk.NORMAL, command=lambda: open_file(filepath))

def open_file(filepath):
    webbrowser.open(filepath)

def reset_inputs():
    name_entry.delete(0, tk.END)
    usn_entry.delete(0, tk.END)
    guide_entry.delete(0, tk.END)
    start_date_entry.delete(0, tk.END)
    college_name_entry.delete(0, tk.END)
    domain_entry.delete(0, tk.END)

def create_pdf():
    global tdate
    name = name_entry.get()
    usn = usn_entry.get()
    guide = guide_entry.get()
    start_date = start_date_entry.get()
    college_name = college_name_entry.get()
    domain = domain_entry.get()
    gender = gender_var.get()
    tdate = datetime.datetime.now().strftime("%B %d, %Y")

    if len(name) == 0:
        messagebox.showerror("Error", "Please enter a name.")
        return
    if len(usn) == 0:
        messagebox.showerror("Error", "Please enter a USN.")
        return
    if len(guide) == 0:
        messagebox.showerror("Error", "Please enter a guide name.")
        return
    if len(start_date) == 0:
        messagebox.showerror("Error", "Please enter a start date.")
        return
    if len(college_name) == 0:
        messagebox.showerror("Error", "Please enter a college name.")
        return
    if len(domain) == 0:
        messagebox.showerror("Error", "Please enter a domain name.")
        return

    create_pdf_from_template(name, usn, guide, start_date, college_name, domain, gender, tdate, "pdf")

def save_as_pdf():
    global tdate
    name = name_entry.get()
    usn = usn_entry.get()
    guide = guide_entry.get()
    start_date = start_date_entry.get()
    college_name = college_name_entry.get()
    domain = domain_entry.get()
    gender = gender_var.get()
    tdate = datetime.datetime.now().strftime("%B %d, %Y")

    if len(name) == 0:
        messagebox.showerror("Error", "Please enter a name.")
        return
    if len(usn) == 0:
        messagebox.showerror("Error", "Please enter a USN.")
        return
    if len(guide) == 0:
        messagebox.showerror("Error", "Please enter a guide name.")
        return
    if len(start_date) == 0:
        messagebox.showerror("Error", "Please enter a start date.")
        return
    if len(college_name) == 0:
        messagebox.showerror("Error", "Please enter a college name.")
        return
    if len(domain) == 0:
        messagebox.showerror("Error", "Please enter a domain name.")
        return

    create_pdf_from_template(name, usn, guide, start_date, college_name, domain, gender, tdate, "pdf")

def save_as_docx():
    global tdate
    name = name_entry.get()
    usn = usn_entry.get()
    guide = guide_entry.get()
    start_date = start_date_entry.get()
    college_name = college_name_entry.get()
    domain = domain_entry.get()
    gender = gender_var.get()
    tdate = datetime.datetime.now().strftime("%B %d, %Y")

    if len(name) == 0:
        messagebox.showerror("Error", "Please enter a name.")
        return
    if len(usn) == 0:
        messagebox.showerror("Error", "Please enter a USN.")
        return
    if len(guide) == 0:
        messagebox.showerror("Error", "Please enter a guide name.")
        return
    if len(start_date) == 0:
        messagebox.showerror("Error", "Please enter a start date.")
        return
    if len(college_name) == 0:
        messagebox.showerror("Error", "Please enter a college name.")
        return
    if len(domain) == 0:
        messagebox.showerror("Error", "Please enter a domain name.")
        return

    create_pdf_from_template(name, usn, guide, start_date, college_name, domain, gender, tdate, "docx")

def send_email():
    email = email_entry.get()
    if not email:
        messagebox.showerror("Error", "Please enter an email address.")
        return

    generated_file = link_label.cget("text")[14:]
    if not os.path.exists(generated_file):
        messagebox.showerror("Error", "The generated file does not exist.")
        return

    # Email configurations
    sender_email = "priyashingri107@gmail.com"  
    sender_password = "fzmfwnhjhjtuaouf"  
    receiver_email = email

    msg = MIMEMultipart()
    msg["From"] = sender_email
    msg["To"] = receiver_email
    msg["Subject"] = "Generated File from HanSoftUmans"

    with open(generated_file, "rb") as attachment:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(attachment.read())
        encoders.encode_base64(part)
        part.add_header(
            "Content-Disposition",
            f"attachment; filename= {generated_file}",
        )
        msg.attach(part)

    try:
        server = smtplib.SMTP("smtp.gmail.com", 587)
        server.starttls()
        server.login(sender_email, sender_password)
        server.sendmail(sender_email, receiver_email, msg.as_string())
        server.quit()
        messagebox.showinfo("Email Sent", "Email sent successfully.")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to send email: {e}")

# main window
window = tk.Tk()
window.title("HanSoftUmans")
window.geometry("400x500") 
window.resizable(True, True)  

# center alignment
screen_width = window.winfo_screenwidth()
screen_height = window.winfo_screenheight()
window_width = 400
window_height = 500
x = (screen_width - window_width) // 2
y = (screen_height - window_height) // 2
window.geometry(f"{window_width}x{window_height}+{x}+{y}")  

# headlinename_entry
headline_label = tk.Label(window, text="HanSoftUmans", font=("Arial", 26, "bold"),fg="blue")
headline_label.grid(row=0, column=0, columnspan=2, pady=10)
subheadline_label = tk.Label(window, text="Internship Certificate", font=("Arial", 16, "bold"),fg="black")
subheadline_label.grid(row=1, column=0, columnspan=2, pady=10)

# labels and entry fields
name_label = tk.Label(window, text="Full Name:")
name_label.grid(row=2, column=0, padx=10, pady=5, sticky=tk.E)
name_entry = tk.Entry(window)
name_entry.grid(row=2, column=1, padx=10, pady=5, sticky=tk.W)
name_entry.focus()

usn_label = tk.Label(window, text="USN:")
usn_label.grid(row=3, column=0, padx=10, pady=5, sticky=tk.E)
usn_entry = tk.Entry(window)
usn_entry.grid(row=3, column=1, padx=10, pady=5, sticky=tk.W)

college_name_label = tk.Label(window, text="College Name:")
college_name_label.grid(row=4, column=0, padx=10, pady=5, sticky=tk.E)
college_name_entry = tk.Entry(window)
college_name_entry.grid(row=4, column=1, padx=10, pady=5, sticky=tk.W)

start_date_label = tk.Label(window, text="Start Date:")
start_date_label.grid(row=5, column=0, padx=10, pady=5, sticky=tk.E)
start_date_entry = DateEntry(window, width=12, background='black', foreground='white', date_pattern='dd/mm/yyyy')
start_date_entry.grid(row=5, column=1, padx=10, pady=5, sticky=tk.W)

guide_label = tk.Label(window, text="Guide:")
guide_label.grid(row=6, column=0, padx=10, pady=5, sticky=tk.E)
guide_entry = tk.Entry(window)
guide_entry.grid(row=6, column=1, padx=10, pady=5, sticky=tk.W)


domain_label = tk.Label(window, text="domain:")
domain_label.grid(row=7, column=0, padx=10, pady=5, sticky=tk.E)
domain_entry = tk.Entry(window)
domain_entry.grid(row=7, column=1, padx=10, pady=5, sticky=tk.W)

# variable
gender_var = tk.StringVar()
gender_var.set("None") 

# radio button
gender_label = tk.Label(window, text="Gender:")
gender_label.grid(row=8, column=0, padx=10, pady=5, sticky=tk.E)
male_radio = tk.Radiobutton(window, text="Male", variable=gender_var, value="Male")
male_radio.grid(row=8, column=1, padx=10, pady=5, sticky=tk.W)
female_radio = tk.Radiobutton(window, text="Female", variable=gender_var, value="Female")
female_radio.grid(row=9, column=1, padx=10, pady=5, sticky=tk.W)


email_label = tk.Label(window, text="Email:")
email_label.grid(row=12, column=0, padx=10, pady=5, sticky=tk.E)
email_entry = tk.Entry(window)
email_entry.grid(row=12, column=1, padx=10, pady=5, sticky=tk.W)




#Saving as PDF and DOCX
save_pdf_button = tk.Button(window, text="Save as PDF", command=save_as_pdf)
save_pdf_button.grid(row=10, column=0, padx=10, pady=10)

save_docx_button = tk.Button(window, text="Save as DOCX", command=save_as_docx)
save_docx_button.grid(row=10, column=1, padx=10, pady=10)

# file name
link_label = tk.Label(window, text="")
link_label.grid(row=10, column=0, columnspan=2, padx=10, pady=5)

open_button = tk.Button(window, text="Open File", state=tk.DISABLED)
open_button.grid(row=11, column=0, columnspan=2, padx=10, pady=5)


# send email button
send_email_button = tk.Button(window, text="Send Email", command=send_email)
send_email_button.grid(row=13, column=0, columnspan=2, padx=10, pady=10)
window.grid_rowconfigure(10, weight=1)  
window.grid_columnconfigure(0, weight=1)  
window.grid_columnconfigure(1, weight=1)  

window.title("HanSoftUmans")

# Increase window size
window_width = 600
window_height = 700
screen_width = window.winfo_screenwidth()
screen_height = window.winfo_screenheight()
x = (screen_width - window_width) // 2
y = (screen_height - window_height) // 2
window.geometry(f"{window_width}x{window_height}+{x}+{y}")
window.mainloop()

